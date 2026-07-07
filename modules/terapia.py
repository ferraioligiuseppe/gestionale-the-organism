# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════════════╗
║  TERAPIA — Percorsi terapeutici PNEV (contenitore)                   ║
║                                                                      ║
║  Un contenitore unico per i percorsi terapeutici dello studio:       ║
║  Vision Therapy, MAPS, Stanza del sale, Osteopatia, Terapia          ║
║  miofunzionale, Sports Vision. Per ogni percorso:                    ║
║    • 📅 Diario sedute (con parte economica → confluisce in incassi)  ║
║    • 🎯 Obiettivi & monitoraggio (esiti → Apprendimento PNEV)        ║
║    • 📄 Relazione AI (carta intestata PNEV)                          ║
║                                                                      ║
║  Tutto entra nel Quadro storico e nell'Assistente PNEV.              ║
╚══════════════════════════════════════════════════════════════════════╝
"""

import json
import datetime
import io
import streamlit as st

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_scheda_vuota(terapia):
    """Scheda terapia VUOTA in Word (.docx), da stampare e compilare a mano."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    h = doc.add_heading(f"Scheda di terapia — {terapia} · Metodo PNEV", level=1)
    doc.add_paragraph("Studio The Organism · Dott. Giuseppe Ferraioli").italic = True
    doc.add_paragraph("")
    for campo in ["Paziente:", "Data:", "N° seduta:", "Professionista:"]:
        doc.add_paragraph(f"{campo} ______________________________")
    doc.add_paragraph("Obiettivo della seduta:")
    doc.add_paragraph("______________________________________________________")
    doc.add_paragraph("Attività svolte:")
    doc.add_paragraph("______________________________________________________")
    doc.add_paragraph("Risposta del paziente: ______________________________")
    for titolo in ["🏥 Procedure svolte IN STUDIO", "🏠 Procedure da fare A CASA"]:
        doc.add_heading(titolo, level=2)
        for i in range(1, 9):
            doc.add_paragraph(f"{i}. ______________________________________________")
    doc.add_paragraph("")
    doc.add_paragraph("Listino €: __________   Sconto €: __________   Incassato €: __________")
    doc.add_paragraph("Note:")
    doc.add_paragraph("______________________________________________________")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_procedure_casa(nome_paz, protocollo, settimana, casa):
    """Elenco procedure A CASA in Word (.docx)."""
    import re as _re
    from docx import Document
    def _pulisci(p):
        p = _re.sub(r'^S\d+\s*·\s*', '', p)
        p = _re.sub(r'\s*\([^)]*\)\s*$', '', p)
        return p.strip()
    doc = Document()
    doc.add_heading("Esercizi da fare a casa — Metodo PNEV", level=1)
    oggi = datetime.date.today().strftime("%d/%m/%Y")
    doc.add_paragraph(f"Studio The Organism · Dott. Giuseppe Ferraioli · {oggi}").italic = True
    doc.add_paragraph(f"Paziente: {nome_paz or '________________'}")
    doc.add_paragraph(f"Percorso: {protocollo} — settimana {settimana}")
    doc.add_paragraph("")
    for p in casa:
        doc.add_paragraph(_pulisci(p), style="List Number")
    doc.add_paragraph("")
    nota = doc.add_paragraph(
        "La spiegazione dettagliata di ogni esercizio è disponibile sul sito pnev.it. "
        "Il lavoro quotidiano a casa è essenziale per i risultati.")
    nota.italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


TERAPIE = ["Vision Therapy", "MAPS", "Stanza del sale", "Osteopatia",
           "Terapia miofunzionale", "Sports Vision"]
RISPOSTA = ["—", "🟢 Buona", "🟡 Parziale", "🔴 Scarsa"]
STATO_OB = ["🟦 In corso", "🟢 Raggiunto", "🟡 Parziale", "⏸️ Sospeso"]
METODI = ["—", "Contanti", "POS / Carta", "Bonifico", "Assegno", "Altro"]


# ── Tabelle ───────────────────────────────────────────────────────────

def _assicura_tabelle(conn):
    try:
        cur = conn.cursor()
        cur.execute("""CREATE TABLE IF NOT EXISTS terapia_sedute(
            id BIGSERIAL PRIMARY KEY, paziente_id BIGINT, terapia TEXT,
            data_seduta DATE, numero INT, professionista TEXT,
            obiettivo TEXT, attivita TEXT, risposta TEXT,
            costo REAL, sconto REAL, incassato REAL, metodo TEXT,
            note TEXT, creato TIMESTAMP DEFAULT NOW());""")
        cur.execute("""CREATE TABLE IF NOT EXISTS terapia_obiettivi(
            id BIGSERIAL PRIMARY KEY, paziente_id BIGINT, terapia TEXT,
            descrizione TEXT, baseline INT, attuale INT, target INT,
            stato TEXT, data_inizio DATE, data_rivalut DATE,
            note TEXT, creato TIMESTAMP DEFAULT NOW());""")
        conn.commit()
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass


def render_terapia(conn=None, paz_id=None, paziente=None):
    st.header("🧘 Percorsi terapeutici")
    st.caption("Vision Therapy, MAPS, Stanza del sale, Osteopatia, Miofunzionale, "
               "Sports Vision. Diario, obiettivi e relazione — tutto in cartella.")

    if conn is None:
        st.info("Connessione non disponibile.")
        return

    # Header con paziente attivo + bottone per cercare/selezionare dall'elenco
    try:
        from .paziente_attivo import header_paziente_attivo
        sel = header_paziente_attivo(conn)
        if sel:
            paz_id = sel
    except Exception:
        pass

    if not paz_id:
        st.info("Seleziona un paziente qui sopra per gestire i suoi percorsi.")
        return

    _assicura_tabelle(conn)

    terapia = st.selectbox("Percorso terapeutico", TERAPIE, key="ter_tipo")
    modo = st.radio("Sezione", ["📅 Diario sedute", "🎯 Obiettivi & monitoraggio",
                                "📄 Relazione PDF"],
                    horizontal=True, key="ter_modo")

    if modo == "📅 Diario sedute":
        _render_diario(conn, paz_id, terapia)
    elif modo == "🎯 Obiettivi & monitoraggio":
        _render_obiettivi(conn, paz_id, terapia)
    else:
        _render_relazione(conn, paz_id, paziente, terapia)


# ── Diario sedute ─────────────────────────────────────────────────────

def _render_diario(conn, paz_id, terapia):
    st.caption(f"Sedute di **{terapia}**. La parte economica confluisce negli incassi.")

    _nome_paz = ""
    try:
        from .quadro_storico import carica_paziente
        _pp = carica_paziente(conn, paz_id)
        if _pp:
            _nome_paz = f"{_pp.get('Cognome','')} {_pp.get('Nome','')}".strip()
    except Exception:
        pass

    st.download_button(
        "🖨️ Scheda terapia VUOTA (Word, da compilare a mano)",
        data=_docx_scheda_vuota(terapia),
        file_name=f"scheda_terapia_vuota_{terapia}.docx".replace(" ", "_"),
        mime=_DOCX_MIME, key="ter_sd_vuota")

    try:
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM terapia_sedute WHERE paziente_id=%s AND terapia=%s",
                    (paz_id, terapia))
        n_fatte = cur.fetchone()[0] or 0
    except Exception:
        n_fatte = 0
        try:
            conn.rollback()
        except Exception:
            pass

    with st.expander("➕ Nuova seduta", expanded=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            data_s = st.date_input("Data", value=datetime.date.today(), key="ter_sd_data")
        with c2:
            numero = st.number_input("N° seduta", min_value=1, step=1,
                                     value=int(n_fatte) + 1, key="ter_sd_num")
        with c3:
            prof = st.text_input("Professionista", key="ter_sd_prof")
        obiettivo = st.text_input("Obiettivo della seduta", key="ter_sd_ob")
        attivita = st.text_area("Attività svolte", height=80, key="ter_sd_att")
        risposta = st.selectbox("Risposta del paziente", RISPOSTA, key="ter_sd_risp")

        st.markdown("**💶 Incasso**")
        e1, e2, e3, e4 = st.columns(4)
        with e1:
            costo = st.number_input("Listino €", min_value=0.0, step=5.0, key="ter_sd_costo")
        with e2:
            sconto = st.number_input("Sconto €", min_value=0.0, step=5.0, key="ter_sd_sconto")
        with e3:
            incassato = st.number_input("Incassato €", min_value=0.0, step=5.0, key="ter_sd_inc")
        with e4:
            metodo = st.selectbox("Metodo", METODI, key="ter_sd_met")
        note = st.text_area("Note", height=70, key="ter_sd_note")

        # ── Procedure assegnate in questa seduta (studio / casa) ──────
        st.markdown("---")
        sel_studio_all, sel_casa_all = _blocco_programma_settimana(conn, paz_id)

        # ── Azioni: stampa · invia su pnev.it · salva ─────────────────
        st.markdown("---")
        a1, a2, a3 = st.columns(3)
        with a1:
            if sel_casa_all:
                st.download_button(
                    "🖨️ Stampa elenco A CASA (Word)",
                    data=_docx_procedure_casa(_nome_paz, terapia, numero, sel_casa_all),
                    file_name=f"procedure_casa_{terapia}.docx".replace(" ", "_"),
                    mime=_DOCX_MIME, key="ter_sd_stampa", use_container_width=True)
        with a2:
            if st.button("📲 Invia su pnev.it", key="ter_sd_invia",
                         disabled=not sel_casa_all, use_container_width=True):
                if _invia_pnev(conn, paz_id, terapia, numero, sel_casa_all):
                    st.success("Procedure di casa inviate alla piattaforma pnev.it.")
                else:
                    st.error("Invio non riuscito.")
        with a3:
            if st.button("💾 Salva seduta", type="primary", key="ter_sd_salva",
                         use_container_width=True):
                ok = _salva_seduta(conn, paz_id, terapia, data_s, numero, prof,
                                   obiettivo, attivita, risposta, costo, sconto,
                                   incassato, metodo, note, sel_studio_all, sel_casa_all)
                if ok:
                    st.success(f"Seduta n° {numero} salvata "
                               f"({len(sel_studio_all)} studio, {len(sel_casa_all)} casa).")
                    st.rerun()
                else:
                    st.error("Salvataggio non riuscito.")

    st.markdown(f"#### Sedute di {terapia} ({n_fatte})")
    _elenco_sedute(conn, paz_id, terapia)


def _scheda_vuota_html(terapia):
    """Scheda terapia VUOTA da stampare e compilare a mano (senza computer)."""
    import datetime as _dt
    riga = "<div style='border-bottom:1px solid #aaa;height:26px'></div>"
    righe_proc = "".join(
        f"<tr><td style='width:30px;text-align:center'>{i}</td>"
        f"<td style='border-bottom:1px solid #aaa'></td>"
        f"<td style='width:60px;border-bottom:1px solid #aaa'></td></tr>"
        for i in range(1, 9))
    return f"""<!doctype html><html lang="it"><head><meta charset="utf-8">
<title>Scheda terapia — {terapia}</title><style>
@page{{size:A4;margin:16mm}} body{{font-family:Georgia,serif;color:#1a1a1a;line-height:1.5}}
h1{{font-size:19px;margin:0 0 2px}} .sub{{color:#555;font-size:12px;margin-bottom:14px}}
table{{border-collapse:collapse;width:100%;margin:6px 0}}
.f{{margin:10px 0}} .f b{{font-size:13px}}
.box{{border:1px solid #aaa;min-height:60px;margin-top:4px}}
.row{{display:flex;gap:16px}} .row>div{{flex:1}}
th{{font-size:12px;text-align:left;border-bottom:2px solid #333;padding:4px}}
td{{padding:6px 4px}} h2{{font-size:14px;margin:14px 0 4px}}
</style></head><body>
<h1>Scheda di terapia — {terapia} · Metodo PNEV</h1>
<div class="sub">Studio The Organism · Dott. Giuseppe Ferraioli</div>
<div class="row">
  <div class="f"><b>Paziente:</b> {riga}</div>
  <div class="f"><b>Data:</b> {riga}</div>
</div>
<div class="row">
  <div class="f"><b>N° seduta:</b> {riga}</div>
  <div class="f"><b>Professionista:</b> {riga}</div>
</div>
<div class="f"><b>Obiettivo della seduta:</b><div class="box"></div></div>
<div class="f"><b>Attività svolte:</b><div class="box"></div></div>
<div class="f"><b>Risposta del paziente:</b> {riga}</div>
<h2>🏥 Procedure svolte IN STUDIO</h2>
<table>{righe_proc}</table>
<h2>🏠 Procedure da fare A CASA</h2>
<table>{righe_proc}</table>
<div class="row">
  <div class="f"><b>Listino €:</b> {riga}</div>
  <div class="f"><b>Sconto €:</b> {riga}</div>
  <div class="f"><b>Incassato €:</b> {riga}</div>
</div>
<div class="f"><b>Note:</b><div class="box"></div></div>
</body></html>"""


def _invia_pnev(conn, paz_id, terapia, numero, casa) -> bool:
    """Salva le procedure di casa come 'inviate' su pnev.it (la pagina di casa
    le leggerà quando la attiviamo)."""
    import json as _json
    try:
        cur = conn.cursor()
        cur.execute("""CREATE TABLE IF NOT EXISTS programma_casa(
            id BIGSERIAL PRIMARY KEY, paziente_id BIGINT,
            protocollo TEXT, settimana INT, procedure JSONB,
            data_assegnazione DATE DEFAULT CURRENT_DATE,
            inviato BOOLEAN DEFAULT FALSE, creato TIMESTAMP DEFAULT NOW());""")
        cur.execute("ALTER TABLE programma_casa ADD COLUMN IF NOT EXISTS tipo TEXT DEFAULT 'casa';")
        cur.execute("""INSERT INTO programma_casa(paziente_id, protocollo, settimana,
            procedure, tipo, inviato) VALUES(%s,%s,%s,%s,'casa',TRUE)""",
            (paz_id, terapia, int(numero), _json.dumps(casa, ensure_ascii=False)))
        conn.commit()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False


def _blocco_programma_settimana(conn, paz_id):
    """Mostra TUTTE le procedure della libreria (per approccio, con ricerca) e
    lascia scegliere liberamente cosa fare IN STUDIO e A CASA in questa seduta.
    I protocolli assegnati restano come scorciatoia (riempimento rapido)."""
    import json as _json
    # carica tutte le procedure attive dalla libreria
    try:
        cur = conn.cursor()
        cur.execute("SELECT approccio, step, nome FROM terapia_libreria "
                    "WHERE attiva=TRUE ORDER BY approccio, step, nome")
        rows = cur.fetchall()
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        rows = []

    if not rows:
        st.caption("La libreria procedure è vuota. Vai in **🧩 Programma PNEV → "
                   "📚 Libreria procedure** per caricarle.")
        return [], []

    # etichetta: [Approccio] (step) Nome
    voci = []
    for appr, step, nome in rows:
        st_lbl = f"{step} · " if step and step != "—" else ""
        voci.append(f"[{appr}] {st_lbl}{nome}")

    st.markdown("**📋 Procedure di questa seduta** — scegli quelle adeguate al paziente:")
    cerca = st.text_input("🔎 Cerca procedura", key="ter_cerca_proc",
                          placeholder="es. saccadi, deglutizione, Marsden…")
    if cerca.strip():
        q = cerca.lower()
        voci_f = [v for v in voci if q in v.lower()]
    else:
        voci_f = voci
    st.caption(f"{len(voci_f)} procedure disponibili"
               + (f" (filtro: «{cerca}»)" if cerca.strip() else ""))

    cstudio, ccasa = st.columns(2)
    with cstudio:
        st.markdown("🏥 **In studio (oggi)**")
        sel_studio = st.multiselect("Procedure in seduta", voci_f, default=[],
                                    key="ter_studio_sel", label_visibility="collapsed")
    with ccasa:
        st.markdown("🏠 **A casa (fino alla prossima)**")
        if st.button("📋 Copia da «In studio»", key="ter_copia_casa"):
            st.session_state["ter_casa_sel"] = list(st.session_state.get("ter_studio_sel", []))
            st.rerun()
        sel_casa = st.multiselect("Procedure a casa", voci_f,
                                  key="ter_casa_sel", label_visibility="collapsed")
    return list(sel_studio), list(sel_casa)


def _foglio_html(nome_paz, protocollo, settimana, casa):
    import datetime as _dt
    import re as _re
    def _pulisci(p):
        p = _re.sub(r'^S\d+\s*·\s*', '', p)
        p = _re.sub(r'\s*\([^)]*\)\s*$', '', p)
        return p.strip()
    def _lista(items):
        if not items:
            return "<p style='color:#888'>—</p>"
        return "<ol>" + "".join(f"<li>{_pulisci(p)}</li>" for p in items) + "</ol>"
    oggi = _dt.date.today().strftime("%d/%m/%Y")
    return f"""<!doctype html><html lang="it"><head><meta charset="utf-8">
<title>Procedure a casa</title><style>
@page{{size:A4;margin:18mm}} body{{font-family:Georgia,serif;color:#1a1a1a;line-height:1.8}}
h1{{font-size:20px;margin:0 0 2px}} .sub{{color:#555;font-size:12px;margin-bottom:14px}}
ol{{margin:10px 0;font-size:15px}} li{{margin:8px 0}}
.foot{{margin-top:26px;font-size:11px;color:#666}}
</style></head><body>
<h1>Esercizi da fare a casa — Metodo PNEV</h1>
<div class="sub">Studio The Organism · Dott. Giuseppe Ferraioli · {oggi}</div>
<p><b>Paziente:</b> {nome_paz or '________________'}<br>
<b>Percorso:</b> {protocollo} — settimana {settimana}</p>
{_lista(casa)}
<p class="foot">La spiegazione dettagliata di ogni esercizio è disponibile sul sito
pnev.it. Il lavoro quotidiano a casa è essenziale per i risultati.</p>
</body></html>"""


def _salva_programma(conn, paz_id, protocollo, settimana, procedure, tipo) -> bool:
    import json as _json
    try:
        cur = conn.cursor()
        cur.execute("""INSERT INTO programma_casa(paziente_id, protocollo, settimana,
            procedure, tipo) VALUES(%s,%s,%s,%s,%s)""",
            (paz_id, protocollo, int(settimana),
             _json.dumps(procedure, ensure_ascii=False), tipo))
        conn.commit()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False


def _salva_programma_casa(conn, paz_id, protocollo, settimana, procedure) -> bool:
    import json as _json
    try:
        cur = conn.cursor()
        cur.execute("""INSERT INTO programma_casa(paziente_id, protocollo, settimana, procedure)
            VALUES(%s,%s,%s,%s)""",
            (paz_id, protocollo, int(settimana),
             _json.dumps(procedure, ensure_ascii=False)))
        conn.commit()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False


def _salva_seduta(conn, paz_id, terapia, data_s, numero, prof, ob, att, risp,
                  costo, sconto, incassato, metodo, note,
                  proc_studio=None, proc_casa=None) -> bool:
    import json as _json
    try:
        cur = conn.cursor()
        cur.execute("ALTER TABLE terapia_sedute ADD COLUMN IF NOT EXISTS procedure_studio TEXT;")
        cur.execute("ALTER TABLE terapia_sedute ADD COLUMN IF NOT EXISTS procedure_casa TEXT;")
        cur.execute("""INSERT INTO terapia_sedute(paziente_id, terapia, data_seduta,
            numero, professionista, obiettivo, attivita, risposta,
            costo, sconto, incassato, metodo, note, procedure_studio, procedure_casa)
            VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
            (paz_id, terapia, data_s, int(numero), prof, ob, att, risp,
             float(costo or 0), float(sconto or 0), float(incassato or 0), metodo, note,
             _json.dumps(proc_studio or [], ensure_ascii=False),
             _json.dumps(proc_casa or [], ensure_ascii=False)))
        conn.commit()
        # confluenza negli incassi: riga nella tabella Sedute (canonica)
        try:
            cur.execute("""INSERT INTO Sedute(paziente_id, Data_Seduta, Terapia,
                Professionista, Costo, Pagato, Note) VALUES(%s,%s,%s,%s,%s,%s,%s)""",
                (paz_id, data_s.strftime("%Y-%m-%d"), terapia, prof or "",
                 float(costo or 0), int(round(incassato or 0)),
                 f"[{terapia}] {ob or ''}".strip()))
            conn.commit()
        except Exception:
            conn.rollback()  # se Sedute ha schema diverso, non blocco il salvataggio
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False


def _elenco_sedute(conn, paz_id, terapia):
    try:
        cur = conn.cursor()
        cur.execute("""SELECT id, data_seduta, numero, professionista, obiettivo,
            attivita, risposta, costo, sconto, incassato, metodo, note,
            procedure_studio, procedure_casa
            FROM terapia_sedute WHERE paziente_id=%s AND terapia=%s
            ORDER BY data_seduta DESC, numero DESC""", (paz_id, terapia))
        righe = cur.fetchall()
    except Exception:
        righe = []
        try:
            conn.rollback()
        except Exception:
            pass
    if not righe:
        st.caption("Nessuna seduta registrata per questo percorso.")
        return
    for (rid, ds, num, prof, ob, att, risp, costo, sconto, inc, met, note,
         p_studio, p_casa) in righe:
        import json as _json
        try:
            l_studio = _json.loads(p_studio) if p_studio else []
        except Exception:
            l_studio = []
        try:
            l_casa = _json.loads(p_casa) if p_casa else []
        except Exception:
            l_casa = []
        ds_str = ds.strftime("%d/%m/%Y") if ds else ""
        cap = f"Seduta n° {num} — {ds_str}"
        if risp and risp != "—":
            cap += f"  ·  {risp}"
        if inc:
            cap += f"  ·  💶 {inc:.0f}€"
        with st.expander(cap):
            e1, e2, e3 = st.columns(3)
            with e1:
                n_data = st.date_input("Data", value=ds or datetime.date.today(),
                                       key=f"ter_ed_data_{rid}")
            with e2:
                n_num = st.number_input("N° seduta", min_value=1, step=1,
                                        value=int(num or 1), key=f"ter_ed_num_{rid}")
            with e3:
                n_prof = st.text_input("Professionista", value=prof or "",
                                       key=f"ter_ed_prof_{rid}")
            n_ob = st.text_input("Obiettivo", value=ob or "", key=f"ter_ed_ob_{rid}")
            n_att = st.text_area("Attività svolte", value=att or "", height=80,
                                 key=f"ter_ed_att_{rid}")
            r1, r2 = st.columns(2)
            with r1:
                n_risp = st.selectbox("Risposta", RISPOSTA,
                                      index=RISPOSTA.index(risp) if risp in RISPOSTA else 0,
                                      key=f"ter_ed_risp_{rid}")
            with r2:
                n_met = st.selectbox("Metodo pagamento", METODI,
                                     index=METODI.index(met) if met in METODI else 0,
                                     key=f"ter_ed_met_{rid}")
            m1, m2, m3 = st.columns(3)
            with m1:
                n_costo = st.number_input("Listino €", min_value=0.0, step=5.0,
                                          value=float(costo or 0), key=f"ter_ed_costo_{rid}")
            with m2:
                n_sconto = st.number_input("Sconto €", min_value=0.0, step=5.0,
                                           value=float(sconto or 0), key=f"ter_ed_sc_{rid}")
            with m3:
                n_inc = st.number_input("Incassato €", min_value=0.0, step=5.0,
                                        value=float(inc or 0), key=f"ter_ed_inc_{rid}")
            n_note = st.text_area("Note", value=note or "", height=70,
                                  key=f"ter_ed_note_{rid}")
            if l_studio or l_casa:
                st.markdown("**🏥 In studio:** " + (", ".join(l_studio) or "—"))
                st.markdown("**🏠 A casa:** " + (", ".join(l_casa) or "—"))
                if l_casa:
                    st.download_button(
                        "🖨️ Stampa elenco A CASA (Word)",
                        data=_docx_procedure_casa("", terapia, num or 1, l_casa),
                        file_name=f"procedure_casa_seduta_{num}.docx",
                        mime=_DOCX_MIME, key=f"ter_ed_stampa_{rid}")
            b1, b2 = st.columns([1, 1])
            with b1:
                if st.button("💾 Salva modifiche", key=f"ter_ed_save_{rid}", type="primary"):
                    if _aggiorna_seduta(conn, rid, n_data, n_num, n_prof, n_ob, n_att,
                                        n_risp, n_costo, n_sconto, n_inc, n_met, n_note):
                        st.success("Seduta aggiornata.")
                        st.rerun()
                    else:
                        st.error("Aggiornamento non riuscito.")
            with b2:
                if st.button("🗑 Elimina", key=f"ter_sd_del_{rid}"):
                    try:
                        cur = conn.cursor()
                        cur.execute("DELETE FROM terapia_sedute WHERE id=%s", (rid,))
                        conn.commit()
                        st.rerun()
                    except Exception:
                        try:
                            conn.rollback()
                        except Exception:
                            pass


def _aggiorna_seduta(conn, rid, data_s, num, prof, ob, att, risp,
                     costo, sconto, inc, met, note) -> bool:
    try:
        cur = conn.cursor()
        cur.execute("""UPDATE terapia_sedute SET data_seduta=%s, numero=%s,
            professionista=%s, obiettivo=%s, attivita=%s, risposta=%s,
            costo=%s, sconto=%s, incassato=%s, metodo=%s, note=%s WHERE id=%s""",
            (data_s, int(num), prof, ob, att, risp, float(costo or 0),
             float(sconto or 0), float(inc or 0), met, note, rid))
        conn.commit()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False


# ── Obiettivi & monitoraggio ──────────────────────────────────────────

def _render_obiettivi(conn, paz_id, terapia):
    st.caption(f"Obiettivi di **{terapia}** (scala 0–10). Alla chiusura l'esito "
               "confluisce nell'Apprendimento PNEV.")

    with st.expander("➕ Nuovo obiettivo", expanded=True):
        with st.form("ter_ob_new", clear_on_submit=True):
            descr = st.text_input("Obiettivo (osservabile)", key="ter_ob_descr")
            c1, c2, c3 = st.columns(3)
            with c1:
                baseline = st.slider("Livello iniziale", 0, 10, 2, key="ter_ob_base")
            with c2:
                target = st.slider("Target", 0, 10, 8, key="ter_ob_targ")
            with c3:
                data_riv = st.date_input("Rivalutazione",
                                         value=datetime.date.today() + datetime.timedelta(weeks=10),
                                         key="ter_ob_riv")
            if st.form_submit_button("💾 Crea obiettivo", type="primary"):
                if descr.strip():
                    if _salva_obiettivo(conn, paz_id, terapia, descr, baseline, target, data_riv):
                        st.success("Obiettivo creato.")
                        st.rerun()
                    else:
                        st.error("Salvataggio non riuscito.")
                else:
                    st.warning("Scrivi l'obiettivo.")

    st.markdown(f"#### Obiettivi di {terapia}")
    _elenco_obiettivi(conn, paz_id, terapia)


def _salva_obiettivo(conn, paz_id, terapia, descr, baseline, target, data_riv) -> bool:
    try:
        cur = conn.cursor()
        cur.execute("""INSERT INTO terapia_obiettivi(paziente_id, terapia, descrizione,
            baseline, attuale, target, stato, data_inizio, data_rivalut)
            VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
            (paz_id, terapia, descr, int(baseline), int(baseline), int(target),
             "🟦 In corso", datetime.date.today(), data_riv))
        conn.commit()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False


def _elenco_obiettivi(conn, paz_id, terapia):
    try:
        cur = conn.cursor()
        cur.execute("""SELECT id, descrizione, baseline, attuale, target, stato, data_rivalut
            FROM terapia_obiettivi WHERE paziente_id=%s AND terapia=%s
            ORDER BY creato DESC""", (paz_id, terapia))
        righe = cur.fetchall()
    except Exception:
        righe = []
        try:
            conn.rollback()
        except Exception:
            pass
    if not righe:
        st.caption("Nessun obiettivo definito.")
        return
    for rid, descr, base, attuale, target, stato, driv in righe:
        st.markdown(f"**{descr}**")
        rng = max(1, (target or 10) - (base or 0))
        prog = min(1.0, max(0.0, ((attuale or 0) - (base or 0)) / rng))
        st.progress(prog, text=f"{stato}  ·  {attuale}/{target} (partenza {base})")
        c1, c2, c3 = st.columns([2, 2, 1])
        with c1:
            nuovo = st.slider("Livello attuale", 0, 10, int(attuale or 0),
                              key=f"ter_ob_upd_{rid}")
        with c2:
            nuovo_stato = st.selectbox("Stato", STATO_OB,
                                       index=STATO_OB.index(stato) if stato in STATO_OB else 0,
                                       key=f"ter_ob_st_{rid}")
        with c3:
            st.write("")
            st.write("")
            if st.button("💾", key=f"ter_ob_save_{rid}", help="Aggiorna"):
                _aggiorna_obiettivo(conn, rid, nuovo, nuovo_stato, paz_id, descr, terapia)
                st.rerun()
        if driv:
            st.caption(f"Rivalutazione: {driv.strftime('%d/%m/%Y') if hasattr(driv,'strftime') else driv}")
        if st.button("🗑 Elimina", key=f"ter_ob_del_{rid}"):
            try:
                cur = conn.cursor()
                cur.execute("DELETE FROM terapia_obiettivi WHERE id=%s", (rid,))
                conn.commit()
                st.rerun()
            except Exception:
                try:
                    conn.rollback()
                except Exception:
                    pass
        st.markdown("<hr style='margin:6px 0;border:none;border-top:1px solid #eee'>",
                    unsafe_allow_html=True)


def _aggiorna_obiettivo(conn, rid, attuale, stato, paz_id, descr, terapia):
    try:
        cur = conn.cursor()
        cur.execute("UPDATE terapia_obiettivi SET attuale=%s, stato=%s WHERE id=%s",
                    (int(attuale), stato, rid))
        conn.commit()
        if stato in ("🟢 Raggiunto", "🟡 Parziale", "⏸️ Sospeso"):
            esito = {"🟢 Raggiunto": "🟢 Migliorato", "🟡 Parziale": "🟡 Stabile / fermo",
                     "⏸️ Sospeso": "⚪ Non valutabile"}.get(stato, "⚪ Non valutabile")
            try:
                cur.execute("""CREATE TABLE IF NOT EXISTS esiti_pnev(
                    id BIGSERIAL PRIMARY KEY, paziente_id BIGINT,
                    data TIMESTAMP DEFAULT NOW(),
                    intervento TEXT, esito TEXT, note TEXT);""")
                cur.execute("INSERT INTO esiti_pnev(paziente_id, intervento, esito, note) "
                            "VALUES(%s,%s,%s,%s)",
                            (paz_id, f"{terapia}: {descr}", esito, "Da obiettivo terapeutico"))
                conn.commit()
            except Exception:
                conn.rollback()
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass


# ── Relazione AI ──────────────────────────────────────────────────────

def _dati_terapia(conn, paz_id, terapia) -> str:
    parti = []
    try:
        cur = conn.cursor()
        cur.execute("""SELECT numero, data_seduta, obiettivo, risposta FROM terapia_sedute
            WHERE paziente_id=%s AND terapia=%s ORDER BY data_seduta DESC LIMIT 20""",
            (paz_id, terapia))
        sd = cur.fetchall()
        if sd:
            parti.append("SEDUTE:")
            for num, ds, ob, risp in sd:
                parti.append(f"- n°{num} {ds}: {ob or ''} ({risp or ''})")
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
    try:
        cur = conn.cursor()
        cur.execute("""SELECT descrizione, baseline, attuale, target, stato
            FROM terapia_obiettivi WHERE paziente_id=%s AND terapia=%s""", (paz_id, terapia))
        ob = cur.fetchall()
        if ob:
            parti.append("\nOBIETTIVI:")
            for descr, base, att, targ, stato in ob:
                parti.append(f"- {descr}: {stato} {att}/{targ} (da {base})")
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
    return "\n".join(parti).strip()


def _render_relazione(conn, paz_id, paziente, terapia):
    st.caption(f"Bozza di relazione del percorso **{terapia}**, con carta intestata "
               "e firma PNEV. L'AI scrive, tu correggi.")

    try:
        from .ai_estrazione import genera_testo, ai_disponibile
        from .diagnosi_assistita import INTESTAZIONE, FIRMA, _identificativi
    except Exception as e:
        st.error(f"Moduli AI/diagnosi non disponibili: {e}")
        return

    if not isinstance(paziente, dict) or not (paziente.get("Cognome") or paziente.get("Nome")):
        try:
            from .quadro_storico import carica_paziente
            p = carica_paziente(conn, paz_id)
            if p:
                paziente = p
        except Exception:
            pass

    dati = _dati_terapia(conn, paz_id, terapia)
    nome = ""
    if isinstance(paziente, dict):
        nome = f"{paziente.get('Cognome','')} {paziente.get('Nome','')}".strip()

    if not dati:
        st.info("Nessun dato ancora salvato per questo percorso: registra prima "
                "sedute o obiettivi.")

    key = f"ter_rel_{paz_id}_{terapia}"
    disabled = not (ai_disponibile() and dati)
    if st.button("🤖 Genera bozza relazione", type="primary", disabled=disabled):
        ident = _identificativi(paziente)
        schema = (
            f"Redigi una bozza di RELAZIONE del percorso terapeutico «{terapia}» "
            "secondo il Metodo PNEV, sui dati qui sotto. Sezioni:\n"
            "1. Dati identificativi\n2. Percorso e finalità\n"
            "3. Andamento delle sedute\n4. Obiettivi e avanzamento\n"
            "5. Considerazioni e indicazioni\n\n"
            "NON scrivere intestazione né firma. Attieniti ai dati; dove mancano, "
            "scrivi «da approfondire».\n\n"
            f"=== DATI IDENTIFICATIVI ===\n{ident}\n\n=== DATI {terapia.upper()} ===\n")
        sistema = ("Sei un terapeuta dello Studio The Organism (Metodo PNEV). "
                   "Italiano, registro clinico, terza persona. NON inventare dati.")
        with st.spinner("L'AI sta scrivendo la relazione…"):
            corpo = genera_testo(schema + dati, sistema=sistema)
        if corpo.startswith("⚠️"):
            st.session_state[key] = corpo
        else:
            st.session_state[key] = INTESTAZIONE + "\n\n" + corpo.strip() + "\n\n" + FIRMA
    if not ai_disponibile():
        st.caption("AI non configurata: la relazione automatica richiede la chiave nei Secrets.")

    testo = st.text_area("Relazione (modificabile)",
                         value=st.session_state.get(key, ""), height=440,
                         key=f"ter_rel_txt_{paz_id}")
    st.download_button("⬇️ Scarica (.txt)", data=testo or "",
                       file_name=f"relazione_{terapia}_{nome or paz_id}.txt".replace(" ", "_"),
                       mime="text/plain", key=f"ter_rel_dl_{paz_id}")
