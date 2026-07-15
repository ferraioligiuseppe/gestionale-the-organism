# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════════════╗
║  PROTOCOLLI PNEV — percorsi definiti nel tempo (livello 2)           ║
║                                                                      ║
║  Un protocollo è una SEQUENZA predefinita di procedure scandita nel  ║
║  tempo (settimana per settimana o per fasi): es. «Apprendimento —    ║
║  stimolazione visiva 10 settimane». Si ASSEGNA a un paziente: il     ║
║  programma nasce già pronto, poi si personalizza.                    ║
║                                                                      ║
║  Include il CONSENSO INFORMATO / CONTRATTO TERAPEUTICO che la        ║
║  famiglia sottoscrive: composizione, durata, importanza del lavoro   ║
║  a casa, impegno della famiglia, costo.                              ║
╚══════════════════════════════════════════════════════════════════════╝
"""

import json
import datetime
import streamlit as st

try:
    from .protocolli_data import PROTOCOLLI as _SEED_PROT
except Exception:
    _SEED_PROT = []


def _assicura_tabelle(conn):
    try:
        cur = conn.cursor()
        cur.execute("""CREATE TABLE IF NOT EXISTS protocolli_assegnati(
            id BIGSERIAL PRIMARY KEY, paziente_id BIGINT,
            nome TEXT, dominio TEXT, durata TEXT,
            data_inizio DATE, settimana_corrente INT DEFAULT 1,
            struttura JSONB, stato TEXT DEFAULT 'In corso',
            creato TIMESTAMP DEFAULT NOW());""")
        cur.execute("""CREATE TABLE IF NOT EXISTS terapia_consensi(
            id BIGSERIAL PRIMARY KEY, paziente_id BIGINT,
            protocollo TEXT, durata TEXT, costo TEXT,
            testo TEXT, firmatario TEXT, accettato BOOLEAN DEFAULT FALSE,
            data TIMESTAMP DEFAULT NOW());""")
        conn.commit()
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass


def render_protocolli(conn, paz_id, paziente=None):
    st.caption("Percorsi PNEV definiti nel tempo. Scegli un protocollo, firmane "
               "il consenso con la famiglia e assegnalo: il programma nasce già pronto.")
    _assicura_tabelle(conn)

    if not _SEED_PROT:
        st.warning("Nessun protocollo disponibile (manca protocolli_data.py).")
        return

    nomi = [p["nome"] for p in _SEED_PROT]
    st.caption("Puoi comporre il **programma PNEV completo** selezionando più protocolli "
               "insieme (visivo, uditivo, riflessi, miofunzionale…). Nelle sedute potrai "
               "poi aggiustare il tiro con le procedure previste e assegnate.")
    scelti = st.multiselect("Protocolli da includere nel programma PNEV", nomi,
                            key="prot_sel")
    prot_list = [p for p in _SEED_PROT if p["nome"] in scelti]
    if not prot_list:
        st.info("Seleziona uno o più protocolli per comporre il programma.")
        return

    domini = " · ".join(sorted({p.get("dominio", "—") for p in prot_list}))
    tappe_tot = sum(len(p.get("settimane", [])) for p in prot_list)
    c1, c2 = st.columns(2)
    c1.metric("Domini inclusi", domini)
    c2.metric("Tappe totali", tappe_tot)

    # anteprima di ciascun protocollo
    for prot in prot_list:
        with st.expander(f"📅 {prot['nome']} — {prot.get('durata','')}", expanded=False):
            for s in prot.get("settimane", []):
                st.markdown(f"**Tappa {s['sett']} — {s.get('fase','')}**")
                for pr in s.get("procedure", []):
                    st.markdown(f"- {pr}")

    st.markdown("---")
    st.markdown("### 📝 Consenso informato / Contratto terapeutico")
    if not paz_id:
        st.info("Seleziona prima un paziente per firmare il consenso e assegnare.")
        return

    nome_paz = ""
    if not isinstance(paziente, dict) or not (paziente.get("Cognome") or paziente.get("Nome")):
        try:
            from .quadro_storico import carica_paziente
            p = carica_paziente(conn, paz_id)
            if p:
                paziente = p
        except Exception:
            pass
    if isinstance(paziente, dict):
        nome_paz = f"{paziente.get('Cognome','')} {paziente.get('Nome','')}".strip()

    # Rilevamento automatico adulto/minore dalla data di nascita (correggibile)
    eta = _eta_paziente(paziente)
    if eta is not None:
        default_idx = 0 if eta >= 18 else 1
        st.caption(f"Età rilevata: **{eta} anni** → "
                   + ("adulto" if eta >= 18 else "minore"))
    else:
        default_idx = 1
        st.caption("Età non disponibile dall'anagrafica: scegli sotto il tipo di consenso.")
    tipo = st.radio("Tipo di consenso", ["🧑 Adulto (firma in proprio)",
                                        "👦 Minore (firma genitore/tutore)"],
                    index=default_idx, horizontal=True, key="prot_tipo")
    is_adult = tipo.startswith("🧑")

    cc1, cc2 = st.columns(2)
    with cc1:
        costo = st.text_input("Costo del percorso (€)", key="prot_costo",
                              placeholder="es. 1.200 € · 120 €/seduta")
        data_inizio = st.date_input("Data inizio", value=datetime.date.today(),
                                    key="prot_inizio")
    with cc2:
        firmatario = st.text_input(
            "Chi sottoscrive" + (" (il paziente)" if is_adult else " (genitore/tutore)"),
            key="prot_firma",
            placeholder=("Nome e cognome del paziente" if is_adult
                         else "Nome e cognome del genitore"))

    testo = _testo_consenso_multi(prot_list, nome_paz, costo, firmatario, is_adult)
    st.text_area("Testo del consenso (modificabile)", value=testo, height=320,
                 key="prot_testo")

    st.download_button(
        "🖨️ Scarica consenso (Word, da stampare o inviare per firma)",
        data=_docx_consenso(st.session_state.get("prot_testo", testo)),
        file_name=f"consenso_pnev_{nome_paz or paz_id}.docx".replace(" ", "_"),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="prot_consenso_dl")

    accetta = st.checkbox(
        ("Il paziente ha letto, compreso e **sottoscrive** il presente consenso, "
         "impegnandosi al lavoro a casa secondo il programma.") if is_adult else
        ("La famiglia ha letto, compreso e **sottoscrive** il presente consenso, "
         "impegnandosi al lavoro a casa secondo il programma."), key="prot_acc")

    if st.button("✅ Firma consenso e assegna programma", type="primary",
                 disabled=not (accetta and firmatario.strip())):
        nomi_sel = ", ".join(p["nome"] for p in prot_list)
        prot_consenso = {"nome": nomi_sel, "durata": "programma combinato",
                         "dominio": ", ".join(sorted({p.get("dominio","—") for p in prot_list})),
                         "settimane": [s for p in prot_list for s in p.get("settimane", [])]}
        ok1 = _salva_consenso(conn, paz_id, prot_consenso, costo,
                              st.session_state.get("prot_testo", testo), firmatario)
        ok2 = all(_assegna_protocollo(conn, paz_id, p, data_inizio) for p in prot_list)
        if ok1 and ok2:
            st.success(f"Programma PNEV assegnato ({len(prot_list)} protocolli). "
                       "Lo trovi nel programma del paziente e nel Quadro storico.")
        else:
            st.warning("Qualcosa non è andato a buon fine nel salvataggio.")

    # consensi/protocolli già presenti
    st.markdown("---")
    _elenco_assegnati(conn, paz_id)


def _eta_paziente(paziente):
    """Età in anni dalla data di nascita del paziente (o None)."""
    if not isinstance(paziente, dict):
        return None
    dn = (paziente.get("Data_Nascita") or paziente.get("data_nascita")
          or paziente.get("Data nascita") or "")
    if not dn:
        return None
    d = dn
    if isinstance(d, str):
        for f in ("%Y-%m-%d", "%d/%m/%Y", "%Y/%m/%d"):
            try:
                d = datetime.datetime.strptime(d[:10], f).date()
                break
            except Exception:
                continue
        if isinstance(d, str):
            return None
    if isinstance(d, datetime.datetime):
        d = d.date()
    try:
        o = datetime.date.today()
        return o.year - d.year - ((o.month, o.day) < (d.month, d.day))
    except Exception:
        return None


def _testo_consenso_multi(prot_list, nome_paz, costo, firmatario, is_adult=False):
    """Consenso per un PROGRAMMA composto da più protocolli (mix PNEV)."""
    elenco = "\n".join(
        f"   • {p.get('nome','')} ({p.get('dominio','—')}, {p.get('durata','—')})"
        for p in prot_list)
    combinato = {
        "nome": "Programma PNEV integrato",
        "dominio": ", ".join(sorted({p.get("dominio", "—") for p in prot_list})),
        "durata": "programma combinato (vedi singoli protocolli)",
        "settimane": [s for p in prot_list for s in p.get("settimane", [])],
    }
    base = _testo_consenso(combinato, nome_paz, costo, firmatario, is_adult)
    ins = ("Il programma proposto integra più approcci PNEV:\n" + elenco + "\n\n")
    return base.replace("1. COMPOSIZIONE DEL PERCORSO\n",
                        ins + "1. COMPOSIZIONE DEL PERCORSO\n")


def _testo_consenso(prot, nome_paz, costo, firmatario, is_adult=False):
    proc_tot = sum(len(s.get("procedure", [])) for s in prot.get("settimane", []))
    intro = (
        "CONSENSO INFORMATO E IMPEGNO TERAPEUTICO — Metodo PNEV\n"
        "Studio The Organism · Dott. Giuseppe Ferraioli\n"
        "────────────────────────────────────────────\n\n"
        f"Paziente: {nome_paz or '________________'}\n"
        f"Protocollo proposto: {prot.get('nome','')}\n"
        f"Dominio: {prot.get('dominio','')}  ·  Durata prevista: {prot.get('durata','')}\n"
        f"Costo del percorso: {costo or '________________'}\n\n"
        "1. COMPOSIZIONE DEL PERCORSO\n"
        f"Il percorso si articola in {len(prot.get('settimane', []))} tappe, per un totale di "
        f"circa {proc_tot} procedure, che evolvono progressivamente (per la parte visiva: "
        "monoculare → bioculare → binoculare) integrando aspetti visivi, percettivi e motori.\n\n"
        "2. IMPORTANZA DEL LAVORO A CASA\n"
        "Il metodo PNEV richiede un allenamento quotidiano a casa. I risultati dipendono in "
        "modo determinante dalla costanza con cui le procedure assegnate vengono svolte "
        "tra una seduta e l'altra. Senza il lavoro a casa il percorso NON può dare i "
        "risultati attesi.\n\n")
    if is_adult:
        impegno = (
            "3. IMPEGNO DEL PAZIENTE\n"
            "Il paziente si impegna a: svolgere quotidianamente gli esercizi assegnati, "
            "rispettare la frequenza delle sedute, comunicare difficoltà o variazioni e "
            "portare avanti il percorso con continuità per l'intera durata prevista.\n\n")
        sottoscr = (
            "5. ACCETTAZIONE E SOTTOSCRIZIONE\n"
            f"Il/la sottoscritto/a {firmatario or nome_paz or '________________'}, in qualità "
            "di paziente, dichiara di aver letto e compreso quanto sopra e di accettarlo in "
            "ogni sua parte, impegnandosi a rispettarlo come un contratto.\n\n"
            "Data e luogo: ____________________     Firma del paziente: ____________________")
    else:
        impegno = (
            "3. IMPEGNO DELLA FAMIGLIA\n"
            "La famiglia si impegna a: garantire lo svolgimento quotidiano degli esercizi, "
            "rispettare la frequenza delle sedute, comunicare difficoltà o variazioni, e "
            "accompagnare il minore con continuità per l'intera durata del percorso.\n\n")
        sottoscr = (
            "5. ACCETTAZIONE E SOTTOSCRIZIONE\n"
            f"Il/la sottoscritto/a {firmatario or '________________'}, in qualità di "
            "genitore/tutore del minore, dichiara di aver letto e compreso quanto sopra e di "
            "accettarlo in ogni sua parte, impegnandosi a rispettarlo come un contratto.\n\n"
            "Data e luogo: ____________________     Firma del genitore/tutore: ____________________")
    natura = (
        "4. NATURA DELL'INTERVENTO\n"
        "L'intervento è di tipo funzionale/riabilitativo e non sostituisce diagnosi o "
        "terapie mediche. I tempi e gli esiti sono individuali e non garantibili a priori.\n\n")
    return intro + impegno + natura + sottoscr


def _salva_consenso(conn, paz_id, prot, costo, testo, firmatario) -> bool:
    try:
        cur = conn.cursor()
        cur.execute("""INSERT INTO terapia_consensi(paziente_id, protocollo, durata,
            costo, testo, firmatario, accettato) VALUES(%s,%s,%s,%s,%s,%s,TRUE)""",
            (paz_id, prot.get("nome", ""), prot.get("durata", ""), costo or "",
             testo, firmatario))
        conn.commit()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False


def _assegna_protocollo(conn, paz_id, prot, data_inizio) -> bool:
    try:
        cur = conn.cursor()
        cur.execute("""INSERT INTO protocolli_assegnati(paziente_id, nome, dominio,
            durata, data_inizio, settimana_corrente, struttura)
            VALUES(%s,%s,%s,%s,%s,1,%s)""",
            (paz_id, prot.get("nome", ""), prot.get("dominio", ""),
             prot.get("durata", ""), data_inizio,
             json.dumps(prot.get("settimane", []), ensure_ascii=False)))
        conn.commit()
        return True
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        return False


def _elenco_assegnati(conn, paz_id):
    try:
        cur = conn.cursor()
        cur.execute("""SELECT id, nome, durata, data_inizio, settimana_corrente, struttura, stato
            FROM protocolli_assegnati WHERE paziente_id=%s ORDER BY creato DESC""", (paz_id,))
        righe = cur.fetchall()
    except Exception:
        righe = []
        try:
            conn.rollback()
        except Exception:
            pass
    if not righe:
        st.caption("Nessun protocollo ancora assegnato a questo paziente.")
        return
    st.markdown("#### Protocolli assegnati")
    st.caption("Qui sotto trovi i protocolli che hai inserito per questo paziente, "
               "con l'avanzamento settimanale.")
    for rid, nome, durata, dini, sett_cur, strut, stato in righe:
        settimane = strut if isinstance(strut, list) else (json.loads(strut) if strut else [])
        tot = len(settimane)
        di = dini.strftime("%d/%m/%Y") if hasattr(dini, "strftime") else str(dini or "")
        st.markdown(f"**{nome}** · {stato} — inizio {di}")
        st.progress(min(1.0, (sett_cur or 1) / tot if tot else 0),
                    text=f"Settimana {sett_cur} di {tot}")
        cc = st.columns([1, 1, 2])
        with cc[0]:
            if st.button("◀ Indietro", key=f"prot_prev_{rid}", disabled=(sett_cur or 1) <= 1):
                _set_settimana(conn, rid, (sett_cur or 1) - 1); st.rerun()
        with cc[1]:
            if st.button("Avanti ▶", key=f"prot_next_{rid}", disabled=(sett_cur or 1) >= tot):
                _set_settimana(conn, rid, (sett_cur or 1) + 1); st.rerun()
        # mostra la settimana corrente
        cur_s = next((s for s in settimane if s.get("sett") == (sett_cur or 1)), None)
        if cur_s:
            st.caption(f"In corso — Settimana {cur_s['sett']}: {cur_s.get('fase','')}")
            for pr in cur_s.get("procedure", []):
                st.markdown(f"- {pr}")
        if st.button("🗑 Rimuovi protocollo", key=f"prot_del_{rid}"):
            try:
                cur = conn.cursor()
                cur.execute("DELETE FROM protocolli_assegnati WHERE id=%s", (rid,))
                conn.commit(); st.rerun()
            except Exception:
                try:
                    conn.rollback()
                except Exception:
                    pass
        st.markdown("<hr style='margin:6px 0;border:none;border-top:1px solid #eee'>",
                    unsafe_allow_html=True)


def _docx_consenso(testo: str) -> bytes:
    """Consenso informato in Word, pronto da stampare o inviare per firma."""
    import io
    from docx import Document
    doc = Document()
    for para in (testo or "").split("\n"):
        doc.add_paragraph(para)
    doc.add_paragraph("")
    doc.add_paragraph("Firma: ______________________________     Data: __________")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_settimana(conn, rid, n):
    try:
        cur = conn.cursor()
        cur.execute("UPDATE protocolli_assegnati SET settimana_corrente=%s WHERE id=%s",
                    (int(n), rid))
        conn.commit()
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
