# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════════════╗
║  MODULISTICA PNEV — schede vuote da stampare (Word .docx)            ║
║                                                                      ║
║  Tutte le schede cliniche in BIANCO, pronte da stampare e compilare  ║
║  a mano quando non si usa il computer. Organizzate per fase della    ║
║  visita. Stessi campi del gestionale → cartaceo e digitale identici. ║
╚══════════════════════════════════════════════════════════════════════╝
"""

import io
import datetime
import streamlit as st

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_LINEA = "______________________________"
_BOX = "______________________________________________________"

# Struttura: area → [ (titolo scheda, [sezioni]) ]
# Ogni sezione è ("Titolo sezione", [campi riga], n_righe_libere)
SCHEDE = {
    "🟢 Accoglienza": [
        ("Scheda anagrafica paziente", [
            ("Dati anagrafici",
             ["Cognome:", "Nome:", "Data di nascita:", "Luogo di nascita:",
              "Codice fiscale:", "Indirizzo:", "Città / CAP:", "Telefono:",
              "Email:", "Scuola / classe (se minore):"], 0),
            ("Genitore / tutore (se minore)",
             ["Cognome e nome:", "Telefono:", "Email:"], 0),
            ("Inviante / motivo", ["Inviante:", "Motivo della visita:"], 2),
        ]),
        ("Consenso al trattamento dei dati (Privacy)", [
            ("Informativa",
             ["Il sottoscritto / la sottoscritta:",
              "in qualità di:  ☐ paziente   ☐ genitore/tutore del minore:"], 0),
            ("Dichiarazione",
             ["Dichiara di aver ricevuto l'informativa sul trattamento dei dati "
              "personali e particolari (art. 13 GDPR) e presta il consenso al "
              "trattamento per le finalità cliniche dello Studio The Organism."], 0),
            ("Firma", ["Data:", "Firma:"], 0),
        ]),
        ("Consenso informato terapeutico", [
            ("Percorso proposto", ["Protocollo / percorso:", "Durata prevista:",
                                    "Costo:"], 0),
            ("Impegno",
             ["Il sottoscritto dichiara di aver compreso la composizione del "
              "percorso, l'importanza del lavoro a casa e l'impegno richiesto, "
              "e di accettarlo in ogni sua parte."], 0),
            ("Firma", ["Data:", "Firma del paziente / genitore-tutore:"], 0),
        ]),
    ],
    "🔍 Valutazione": [
        ("Anamnesi The Organism", [
            ("Gravidanza e sviluppo",
             ["Gravidanza e parto:", "Sviluppo motorio:", "Sviluppo linguistico:",
              "Alimentazione:", "Sonno e respiro:"], 0),
            ("Salute", ["Patologie / interventi:", "Familiarità:",
                        "Terapie in corso:"], 2),
            ("Osservazioni", [], 4),
        ]),
        ("Valutazione visuo-percettiva", [
            ("Acuità visiva", ["OD lontano:", "OS lontano:", "OD vicino:",
                               "OS vicino:"], 0),
            ("Refrazione", ["Abituale OD:", "Abituale OS:", "Soggettiva OD:",
                            "Soggettiva OS:"], 0),
            ("Note funzionali", [], 4),
        ]),
        ("DEM — scheda di registrazione", [
            ("Tempi", ["Tempo Verticale (A+B):", "Tempo Orizzontale (C):",
                       "Errori:"], 0),
            ("Esito", ["Rapporto:", "Tipo (I/II/III/IV):", "Età/classe:"], 0),
        ]),
        ("Getman — manipolazione visiva", [
            ("Risposte", ["Triangolo (POV/Capov./Entr.):", "Semisfera:",
                          "T:", "L:"], 0),
            ("Esito", ["Punteggio /12:", "Classe equivalente:"], 0),
        ]),
        ("Groffman — visual tracing", [
            ("Tempi e numeri (A-E)",
             ["A — n°___ sec___", "B — n°___ sec___", "C — n°___ sec___",
              "D — n°___ sec___", "E — n°___ sec___"], 0),
            ("Esito", ["Punteggio totale:", "Età:", "Norma:"], 0),
        ]),
        ("Bilancio uditivo", [
            ("Questionario", [], 6),
            ("Esito", ["Note:", "Invio specialistico:"], 0),
        ]),
        ("INPP — riflessi primitivi", [
            ("Riflessi (0-4)",
             ["Moro:", "TLR:", "ATNR:", "STNR:", "Galant:", "Plantare:"], 0),
            ("Note", [], 3),
        ]),
        ("Logopedia / SMOF", [
            ("Miofunzionale", ["Respirazione:", "Deglutizione:",
                               "Postura linguale:", "Frenulo:"], 0),
            ("Linguaggio / Fluenza", ["Comprensione:", "Fonologia:",
                                       "Fluenza:"], 0),
            ("Note", [], 3),
        ]),
    ],
    "🎧 Terapia": [
        ("Scheda di seduta (terapia)", [
            ("Dati", ["Paziente:", "Data:", "N° seduta:", "Professionista:"], 0),
            ("Seduta", ["Obiettivo:", "Attività svolte:", "Risposta:"], 0),
            ("Procedure IN STUDIO", [], 8),
            ("Procedure A CASA", [], 8),
            ("Incasso", ["Listino €:", "Sconto €:", "Incassato €:", "Metodo:"], 0),
            ("Note", [], 2),
        ]),
        ("Diario degli esercizi a casa (per la famiglia)", [
            ("Settimana", ["Paziente:", "Percorso:", "Settimana:"], 0),
            ("Esercizi (segna ✓ ogni giorno)", [], 10),
        ]),
    ],
    "📝 Relazioni": [
        ("Foglio relazione clinica (intestato, vuoto)", [
            ("Intestazione", ["Paziente:", "Data:"], 0),
            ("Relazione", [], 18),
            ("Firma", ["Dott. Giuseppe Ferraioli — Neuropsicologo · Optometrista"], 0),
        ]),
    ],
}


def _build_docx(titolo, sezioni):
    from docx import Document
    doc = Document()
    doc.add_heading(titolo, level=1)
    intest = doc.add_paragraph("Studio The Organism · Metodo PNEV · "
                               "Dott. Giuseppe Ferraioli")
    intest.italic = True
    doc.add_paragraph("")
    for sez_titolo, campi, n_righe in sezioni:
        if sez_titolo:
            doc.add_heading(sez_titolo, level=2)
        for c in campi:
            doc.add_paragraph(f"{c} {_LINEA}" if c.endswith(":") else c)
        for _ in range(n_righe):
            doc.add_paragraph(_BOX)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_modulistica(conn=None, paz_id=None, paziente=None):
    st.header("📄 Modulistica / Schede da stampare")
    st.caption("Tutte le schede cliniche in bianco, in Word, da stampare e "
               "compilare a mano. Stessi campi del gestionale. Apri in Word e, "
               "se vuoi il PDF, da Word: File → Salva come → PDF.")

    for area, lista in SCHEDE.items():
        st.markdown(f"### {area}")
        for titolo, sezioni in lista:
            c1, c2 = st.columns([3, 1])
            with c1:
                st.markdown(f"**{titolo}**")
            with c2:
                st.download_button(
                    "⬇️ Word", data=_build_docx(titolo, sezioni),
                    file_name=f"{titolo}.docx".replace("/", "-"),
                    mime=_DOCX_MIME, key=f"mod_{titolo}",
                    use_container_width=True)
        st.markdown("<hr style='margin:8px 0;border:none;border-top:1px solid #eee'>",
                    unsafe_allow_html=True)
